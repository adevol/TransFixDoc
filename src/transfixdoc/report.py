"""PDF and Word report generation."""

import html
import re
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer

from transfixdoc.models import Document


def write_report(document: Document, output_path: Path) -> Path:
    """Write a report. Dispatches on file extension (.pdf or .docx).

    Args:
        document: Translated or corrected document.
        output_path: Report path. Suffix determines the format.

    Returns:
        Written report path.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    suffix = output_path.suffix.lower()
    if suffix == ".docx":
        return _write_docx(document, output_path)
    if suffix == ".pdf":
        return _write_pdf(document, output_path)
    raise ValueError(f"Unsupported report format '{suffix}'. Use .pdf or .docx.")


def _write_pdf(document: Document, output_path: Path) -> Path:
    styles = getSampleStyleSheet()
    story = [Paragraph(f"Report for {html.escape(str(document.source_path))}", styles["Title"])]

    for page in document.pages:
        story.extend([Spacer(1, 12), Paragraph(f"Page {page.page_number}", styles["Heading2"])])
        if page.image_path and page.image_path.exists():
            story.extend([Image(str(page.image_path), width=420, height=594), Spacer(1, 12)])
        text = f"<b>Error:</b> {html.escape(page.error)}" if page.error else _markdown_bold(page.text)
        story.append(Paragraph(text.replace("\n", "<br/>"), styles["BodyText"]))
        story.append(PageBreak())

    SimpleDocTemplate(str(output_path), pagesize=A4).build(story)
    return output_path


def _write_docx(document: Document, output_path: Path) -> Path:
    from docx import Document as DocxDocument
    from docx.shared import Inches

    doc = DocxDocument()
    doc.add_heading(f"Report for {document.source_path}", level=1)

    for page in document.pages:
        doc.add_heading(f"Page {page.page_number}", level=2)
        if page.image_path and page.image_path.exists():
            doc.add_picture(str(page.image_path), width=Inches(6))
        if page.error:
            paragraph = doc.add_paragraph()
            paragraph.add_run("Error: ").bold = True
            paragraph.add_run(page.error)
        else:
            for line in page.text.split("\n"):
                paragraph = doc.add_paragraph()
                for chunk, is_bold in _split_bold(line):
                    paragraph.add_run(chunk).bold = is_bold
        doc.add_page_break()

    doc.save(str(output_path))
    return output_path


def _markdown_bold(text: str) -> str:
    escaped = html.escape(text)
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)


def _split_bold(line: str) -> list[tuple[str, bool]]:
    parts = re.split(r"(\*\*.+?\*\*)", line)
    return [
        (part[2:-2], True) if part.startswith("**") and part.endswith("**") else (part, False)
        for part in parts
        if part
    ]
